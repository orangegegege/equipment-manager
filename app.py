import streamlit as st
import pandas as pd
from supabase import create_client, Client
from datetime import datetime, timedelta 
import time
import os
import io 
from fpdf import FPDF 

# Word 處理套件
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ==========================================
# 🎨 [色彩與基本設定]
# ==========================================
NAV_HEIGHT = "80px"
NAV_BG_COLOR = "#E88B00"       # 你的橘色
PAGE_BG_COLOR = "#F5F5F5"      # 淺灰底
LOGO_URL = "https://obmikwclquacitrwzdfc.supabase.co/storage/v1/object/public/logos/logo.png"

# 🔥 統一管理的分類清單
CATEGORY_OPTIONS = ["手工具", "一般器材", "廚具", "清潔用品", "文具用品", "其他"]

# ⚠️ 字體設定 (僅供 PDF 使用)
FONT_FILE = "TaipeiSansTCBeta-Regular.ttf"

# --- 1. Supabase 連線 ---
@st.cache_resource
def init_connection():
    try:
        url = st.secrets["SUPABASE"]["URL"]
        key = st.secrets["SUPABASE"]["KEY"]
        return create_client(url, key)
    except Exception as e:
        st.error(f"Supabase 連線失敗: {e}")
        return None

supabase: Client = init_connection()

# --- 2. 圖片上傳 ---
def upload_image(file):
    if not file: return None
    try:
        bucket_name = st.secrets["SUPABASE"]["BUCKET"]
        file_ext = file.name.split('.')[-1]
        file_name = f"{int(time.time())}_{file.name}"
        supabase.storage.from_(bucket_name).upload(file_name, file.getvalue(), file_options={"content-type": file.type})
        return supabase.storage.from_(bucket_name).get_public_url(file_name)
    except Exception as e:
        st.error(f"上傳失敗: {e}")
        return None

# --- 3. 資料庫 CRUD (新增 borrowed 欄位處理) ---
def load_data():
    # 記得要在 Supabase 執行 SQL: ALTER TABLE equipment ADD COLUMN borrowed INT DEFAULT 0;
    response = supabase.table("equipment").select("*").order("id", desc=True).execute()
    df = pd.DataFrame(response.data)
    # 防呆：如果資料庫還沒加欄位，這裡手動補上避免報錯
    if 'borrowed' not in df.columns and not df.empty:
        df['borrowed'] = 0
    return df

def add_equipment_to_db(data):
    supabase.table("equipment").insert(data).execute()

def update_equipment_in_db(uid, updates):
    supabase.table("equipment").update(updates).eq("uid", uid).execute()

def delete_equipment_from_db(uid):
    supabase.table("equipment").delete().eq("uid", uid).execute()

# --- 輔助函式 ---
def get_taiwan_time_str():
    tw_time = datetime.utcnow() + timedelta(hours=8)
    return tw_time.strftime('%Y-%m-%d %H:%M')

def get_today_str():
    tw_time = datetime.utcnow() + timedelta(hours=8)
    return tw_time.strftime('%Y-%m-%d')

# 🔥🔥🔥 核心邏輯：計算狀態標籤 🔥🔥🔥
def get_status_display(row):
    # 先看有沒有被人工強制設定狀態
    manual_status = row.get('status', '在庫')
    if manual_status in ['維修中', '報廢']:
        return manual_status, "grey"
    
    # 接著看數量邏輯
    total = row.get('quantity', 1)
    borrowed = row.get('borrowed', 0)
    available = total - borrowed
    
    if available <= 0:
        return "🔴 已借完 / 暫無庫存", "red"
    elif borrowed > 0:
        return f"⚠️ 部分在庫 (剩 {available})", "orange"
    else:
        return f"✅ 足額在庫 ({available}/{total})", "green"

# ==========================================
# PDF 與 Word 生成 (沿用之前的邏輯，稍微適應新資料)
# ==========================================
class PDFReport(FPDF):
    def __init__(self):
        super().__init__(orientation='L', unit='mm', format='A4')
        self.set_auto_page_break(auto=False) 

    def header(self):
        if os.path.exists(FONT_FILE):
            try:
                self.add_font('ChineseFont', '', FONT_FILE)
                self.set_font('ChineseFont', '', 12)
            except:
                self.set_font("Helvetica", size=12)
        
        self.set_font_size(24)
        self.cell(0, 15, txt="團隊器材借用 / 清點單", ln=1, align='C')
        self.set_font_size(10)
        self.cell(0, 8, txt=f"製表日期: {get_taiwan_time_str()}", ln=1, align='R')
        self.line(10, self.get_y(), 287, self.get_y())
        self.ln(2)
        
        # 表頭
        self.set_font_size(12)
        self.set_fill_color(232, 139, 0) 
        self.set_text_color(255, 255, 255) 
        self.set_line_width(0.3)
        headers = ["分類項目", "編號", "器材名稱", "借用數量", "營前清點", "離營清點", "營後清點"]
        col_w = [35, 30, 80, 20, 37, 37, 37] 
        for i, h in enumerate(headers):
            self.cell(col_w[i], 10, h, border=1, align='C', fill=True)
        self.ln()
        self.set_text_color(0, 0, 0) 

    def footer(self):
        self.set_y(-25)
        if os.path.exists(FONT_FILE): self.set_font('ChineseFont', '', 12)
        self.cell(90, 10, "器材負責人：__________________", align='L')
        self.cell(90, 10, "活動負責人：__________________", align='C')
        self.cell(90, 10, "指導老師：__________________", align='R')

def create_pdf(cart_data, text_display_map):
    pdf = PDFReport()
    pdf.add_page()
    if os.path.exists(FONT_FILE): pdf.set_font('ChineseFont', '', 11)
    else: pdf.set_font("Helvetica", size=11)

    col_w = [35, 30, 80, 20, 37, 37, 37] 
    total_rows = len(cart_data)
    fill = False 
    pdf.set_fill_color(245, 245, 245)

    for i in range(total_rows):
        if pdf.get_y() > 170:
            pdf.add_page()
            force_new_page_header = True 
        else: force_new_page_header = False

        item = cart_data[i]
        cat = item['category']
        
        draw_top = (i == 0 or cart_data[i-1]['category'] != cat or force_new_page_header)
        draw_bottom = (i == total_rows - 1 or cart_data[i+1]['category'] != cat or (pdf.get_y() + 10 > 170))
        
        cat_border = 'LR' 
        if draw_top: cat_border += 'T'
        if draw_bottom: cat_border += 'B'
        
        cat_display = text_display_map.get(i, "")
        if force_new_page_header: cat_display = cat
        
        pdf.cell(col_w[0], 10, cat_display, border=cat_border, align='C', fill=False)
        pdf.cell(col_w[1], 10, str(item['uid']), border=1, align='C', fill=fill)
        
        name = str(item['name'])
        if pdf.get_string_width(name) > col_w[2] - 2: name = name[:14] + "..."
        pdf.cell(col_w[2], 10, name, border=1, align='C', fill=fill)
        
        # 🔥 這裡是借用數量，不是總庫存
        pdf.cell(col_w[3], 10, str(item['borrow_qty']), border=1, align='C', fill=fill)
        
        pdf.cell(col_w[4], 10, "", border=1, align='C', fill=fill)
        pdf.cell(col_w[5], 10, "", border=1, align='C', fill=fill)
        pdf.cell(col_w[6], 10, "", border=1, align='C', fill=fill)
        pdf.ln()
        fill = not fill 
    return pdf.output()

def set_cell_bg(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_word(cart_data):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.left_margin = Mm(12)
    section.right_margin = Mm(12)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    heading = doc.add_paragraph("團隊器材借用 / 清點單")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.runs[0]
    run.font.size = Pt(24)
    run.font.name = "Microsoft JhengHei"
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
    run.bold = True
    
    date_para = doc.add_paragraph(f"製表日期: {get_taiwan_time_str()}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.autofit = False 
    headers = ["分類項目", "編號", "器材名稱", "借用數量", "營前清點", "離營清點", "營後清點"]
    widths = [12, 10, 30, 8, 13, 13, 13] 
    total_width_mm = 273 
    
    hdr_row = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = text
        set_cell_bg(cell, "E88B00")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = "Microsoft JhengHei"
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
        cell.width = Mm(total_width_mm * widths[i] / 100)

    for idx, item in enumerate(cart_data):
        row_cells = table.add_row().cells
        row_cells[0].text = item['category']
        row_cells[1].text = str(item['uid'])
        row_cells[2].text = str(item['name'])
        row_cells[3].text = str(item['borrow_qty'])
        
        for i, cell in enumerate(row_cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.width = Mm(total_width_mm * widths[i] / 100)
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run()
            run.font.name = "Microsoft JhengHei"
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    col_idx = 0
    start_row = 1 
    while start_row < len(table.rows):
        cat_text = table.rows[start_row].cells[col_idx].text
        end_row = start_row + 1
        while end_row < len(table.rows) and table.rows[end_row].cells[col_idx].text == cat_text:
            table.rows[end_row].cells[col_idx].text = "" 
            end_row += 1
        if end_row > start_row + 1:
            table.rows[start_row].cells[col_idx].merge(table.rows[end_row - 1].cells[col_idx])
        start_row = end_row

    doc.add_paragraph("\n") 
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.autofit = True
    sig_table.width = Mm(273)
    sig_cells = sig_table.rows[0].cells
    sig_cells[0].text = "器材負責人：__________________"
    sig_cells[1].text = "活動負責人：__________________"
    sig_cells[2].text = "指導老師：__________________"
    for cell in sig_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    sig_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- 頁面設定 ---
st.set_page_config(page_title="器材管理系統", layout="wide", page_icon="📦", initial_sidebar_state="collapsed")

# 初始化購物車 (這次改成 Dictionary 來記數量 {uid: qty})
if 'cart' not in st.session_state:
    st.session_state.cart = {} # Key: uid, Value: quantity (預設 1)

# ==========================================
# 🛠️ CSS 樣式表
# ==========================================
st.markdown(f"""
<style>
    header[data-testid="stHeader"] {{ display: none; }}
    .stApp {{ background-color: {PAGE_BG_COLOR} !important; }}
    .main .block-container {{
        padding-top: 100px !important; 
        max-width: 1200px !important;
    }}
    #my-fixed-header {{
        position: fixed;
        top: 0;
        left: 0;
        width: 100%;
        height: {NAV_HEIGHT};
        background-color: {NAV_BG_COLOR};
        z-index: 999999;
        box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        display: flex;
        align-items: center;
        padding-left: 30px;
        padding-right: 30px;
        justify-content: space-between; 
    }}
    div[data-testid="stVerticalBlockBorderWrapper"] {{
        background-color: white !important;
        border: 1px solid #ddd !important;
        border-radius: 8px !important;
        padding: 20px !important;
    }}
    .stButton > button {{
        border-radius: 6px;
        background-color: white;
        color: #333;
        border: 1px solid #ccc;
    }}
    .stButton > button[kind="primary"] {{
        background-color: {NAV_BG_COLOR} !important;
        color: white !important;
        border: none !important;
    }}
    div[data-testid="stPills"] button[aria-selected="true"] {{
        background-color: {NAV_BG_COLOR} !important;
        color: white !important;
        border-color: {NAV_BG_COLOR} !important;
    }}
</style>
""", unsafe_allow_html=True)

# --- 狀態管理 ---
if 'is_admin' not in st.session_state: st.session_state.is_admin = False
if 'current_page' not in st.session_state: st.session_state.current_page = "home"
def go_to(page): st.session_state.current_page = page
def perform_logout(): 
    st.session_state.is_admin = False
    st.session_state.cart = {}
    for key in list(st.session_state.keys()):
        if key.startswith("check_"): del st.session_state[key]
    go_to("home")
def perform_login():
    if st.session_state.password_input == st.secrets["ADMIN_PASSWORD"]:
        st.session_state.is_admin = True
        go_to("home")
    else: st.error("密碼錯誤")

# ==========================================
# 彈窗：檢視清單與借用確認 (核心邏輯升級)
# ==========================================
@st.dialog("📋 借用清單確認", width="large")
def show_cart_modal(df):
    if not st.session_state.cart:
        st.info("清單目前是空的，請先勾選器材！")
        if st.button("關閉"): st.rerun()
    else:
        # 1. 準備資料，並讓使用者輸入數量
        cart_uids = list(st.session_state.cart.keys())
        cart_rows = df[df['uid'].isin(cart_uids)].copy()
        
        # 強制依分類排序
        cart_rows = cart_rows.sort_values(by=['category', 'uid'])
        
        st.info("💡 請確認以下器材與借用數量，按下「確認借用」後將扣除庫存並產生單據。")
        
        # 用來存最終要輸出的資料
        final_borrow_list = []
        
        # 建立編輯介面
        for i, row in cart_rows.iterrows():
            c1, c2, c3, c4 = st.columns([2, 3, 2, 2])
            with c1: st.write(f"**{row['category']}**")
            with c2: st.write(f"{row['name']} (#{row['uid']})")
            with c3:
                # 計算剩餘量 (Total - Borrowed)
                avail = row['quantity'] - row.get('borrowed', 0)
                st.caption(f"庫存剩餘: {avail}")
            with c4:
                # 數量輸入框 (預設 1，最大不能超過剩餘量)
                # 這裡要注意：如果剩餘量是 0，理論上不該能加進來，但防呆一下
                max_val = max(1, avail) 
                borrow_qty = st.number_input(
                    "借用數量", 
                    min_value=1, 
                    max_value=max_val, 
                    value=st.session_state.cart.get(row['uid'], 1), 
                    key=f"qty_{row['uid']}",
                    label_visibility="collapsed"
                )
                
                # 更新 session 裡的數量紀錄
                st.session_state.cart[row['uid']] = borrow_qty
                
                # 準備資料給 PDF/Word
                item_dict = row.to_dict()
                item_dict['borrow_qty'] = borrow_qty
                final_borrow_list.append(item_dict)

        st.markdown("---")
        
        # 3. 格式選擇與確認借用
        col_opt, col_action = st.columns([1, 1])
        with col_opt:
            export_format = st.radio("選擇匯出格式：", ["PDF 文件 (.pdf)", "Word 文件 (.docx)"])
            
        with col_action:
            st.write("") 
            st.write("") 
            
            # 🔥🔥🔥 核心按鈕：確認借用 & 下載 🔥🔥🔥
            # Streamlit 的 download_button 有 callback 功能，我們利用它來更新資料庫
            
            def perform_checkout():
                # 1. 更新資料庫：扣除庫存 (增加 borrowed 數量)
                try:
                    for item in final_borrow_list:
                        current_borrowed = item.get('borrowed', 0)
                        new_borrowed = current_borrowed + item['borrow_qty']
                        update_equipment_in_db(item['uid'], {'borrowed': new_borrowed})
                    
                    st.toast("✅ 借用成功！庫存已扣除。")
                    
                    # 2. 清空購物車
                    st.session_state.cart = {}
                    for key in list(st.session_state.keys()):
                        if key.startswith("check_"): st.session_state[key] = False
                    
                    # 稍微延遲讓 user 看到 toast
                    time.sleep(1)
                except Exception as e:
                    st.error(f"資料庫更新失敗: {e}")

            # 準備檔案資料
            today_date = get_today_str()
            file_prefix = f"equipment_list_{today_date}"
            
            # 計算垂直置中文字位置 (給 PDF 用)
            text_display_map = {} 
            start_index = 0
            total_rows = len(final_borrow_list)
            for i in range(total_rows + 1):
                if i == total_rows or final_borrow_list[i]['category'] != final_borrow_list[start_index]['category']:
                    count = i - start_index
                    center_offset = count // 2
                    center_row = start_index + center_offset
                    text_display_map[center_row] = final_borrow_list[start_index]['category']
                    start_index = i

            if export_format == "PDF 文件 (.pdf)":
                try:
                    pdf_bytes = create_pdf(final_borrow_list, text_display_map)
                    st.download_button(
                        label="🚀 確認借用並下載 (PDF)",
                        data=bytes(pdf_bytes), 
                        file_name=f"{file_prefix}.pdf",
                        mime="application/pdf",
                        type="primary",
                        use_container_width=True,
                        on_click=perform_checkout # 🔥 按下即扣庫存
                    )
                except Exception as e: st.error(f"PDF 錯誤: {e}")
                    
            elif export_format == "Word 文件 (.docx)":
                try:
                    word_bytes = create_word(final_borrow_list)
                    st.download_button(
                        label="🚀 確認借用並下載 (Word)",
                        data=word_bytes,
                        file_name=f"{file_prefix}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True,
                        on_click=perform_checkout # 🔥 按下即扣庫存
                    )
                except Exception as e: st.error(f"Word 錯誤: {e}")

        if st.button("🗑️ 取消 / 清空", use_container_width=True):
            st.session_state.cart = {}
            for key in list(st.session_state.keys()):
                if key.startswith("check_"): st.session_state[key] = False
            st.rerun()

# ==========================================
# Header 組件
# ==========================================
def render_header(df_for_count=None):
    st.markdown(f"""
    <div id="my-fixed-header">
        <div style="display:flex; align-items:center;">
            <img src="{LOGO_URL}" style="height: 50px; object-fit: contain;">
        </div>
    </div>
    """, unsafe_allow_html=True)

# ==========================================
# 彈窗：新增/編輯 (已更新 borrowed 欄位)
# ==========================================
@st.dialog("➕ 新增器材", width="small")
def show_add_modal():
    st.caption("填寫詳細資訊")
    with st.form("add_form", clear_on_submit=True):
        name = st.text_input("名稱", placeholder="例如：活動扳手")
        uid = st.text_input("編號", placeholder="例如：TOOL-001")
        c1, c2 = st.columns(2)
        cat = c1.selectbox("分類", CATEGORY_OPTIONS, index=None, placeholder="--請選擇--")
        status = c2.selectbox("狀態", ["在庫", "維修中", "報廢"], index=0, placeholder="--請選擇--") # 移除借出中，改由數量控制
        
        c3, c4 = st.columns(2)
        qty = c3.number_input("總數量 (Total)", min_value=1, value=1, step=1)
        loc = c4.text_input("位置", value="儲藏室")
        file = st.file_uploader("照片", type=['jpg','png'])
        
        if st.form_submit_button("新增", type="primary", use_container_width=True):
            if name and uid and cat:
                url = upload_image(file) if file else None
                data_payload = {
                    "uid": uid, "name": name, "category": cat, "status": status, 
                    "location": loc, "quantity": qty, "borrowed": 0, # 新增預設為 0
                    "image_url": url, "updated_at": datetime.now().strftime("%Y-%m-%d")
                }
                try:
                    add_equipment_to_db(data_payload)
                    st.toast(f"🎉 成功新增：{name}"); time.sleep(1); st.rerun()
                except Exception as e: st.error(f"寫入失敗: {e}")
            else: st.warning("請填寫完整資訊")

@st.dialog("⚙️ 編輯/管理器材", width="small")
def show_edit_modal(item):
    st.caption(f"正在編輯：{item['name']} (#{item['uid']})")
    if item['image_url']: st.image(item['image_url'], width=100)
    with st.form("edit_form"):
        new_name = st.text_input("名稱", value=item['name'])
        c1, c2 = st.columns(2)
        try: cat_idx = CATEGORY_OPTIONS.index(item['category'])
        except: cat_idx = 0
        new_cat = c1.selectbox("分類", CATEGORY_OPTIONS, index=cat_idx)
        
        # 狀態選單：這裡只留人工強制狀態
        status_opts = ["在庫", "維修中", "報廢"]
        try: status_idx = status_opts.index(item['status'])
        except: status_idx = 0
        new_status = c2.selectbox("特殊狀態 (若無異常請選在庫)", status_opts, index=status_idx)
        
        c3, c4 = st.columns(2)
        new_qty = c3.number_input("總數量 (Total)", min_value=1, value=item.get('quantity', 1), step=1)
        # 🔥 管理員可以手動調整「已借出」數量 (校正用)
        new_borrowed = c4.number_input("已借出 (校正用)", min_value=0, max_value=new_qty, value=item.get('borrowed', 0), step=1)
        
        new_loc = st.text_input("位置", value=item['location'] or "")
        new_file = st.file_uploader("更換照片", type=['jpg','png'])
        
        col_update, col_delete = st.columns([1, 1])
        submitted = col_update.form_submit_button("💾 儲存更新", type="primary", use_container_width=True)
        delete_confirm = col_delete.checkbox("確認刪除此器材")
        
        if submitted:
            if delete_confirm:
                delete_equipment_from_db(item['uid']); st.toast("🗑️ 已刪除"); time.sleep(1); st.rerun()
            else:
                final_url = upload_image(new_file) if new_file else item['image_url']
                updates = {
                    "name": new_name, "category": new_cat, "status": new_status, 
                    "quantity": new_qty, "borrowed": new_borrowed, # 更新借出量
                    "location": new_loc, "image_url": final_url, 
                    "updated_at": datetime.now().strftime("%Y-%m-%d")
                }
                update_equipment_in_db(item['uid'], updates); st.toast("✅ 更新成功"); time.sleep(1); st.rerun()

# ==========================================
# 主頁面
# ==========================================
def main_page():
    render_header()
    
    st.markdown("""
        <style>
        .header-buttons {
            position: fixed; top: 20px; right: 30px; z-index: 9999999;
        }
        </style>
    """, unsafe_allow_html=True)
    
    df = load_data()
    
    with st.container():
        st.markdown('<div class="header-buttons">', unsafe_allow_html=True)
        if not st.session_state.is_admin:
            cart_count = len(st.session_state.cart)
            if st.button(f"📋 借用清單 ({cart_count})", type="primary"):
                show_cart_modal(df)
        st.markdown('</div>', unsafe_allow_html=True)

    c_title, c_actions = st.columns([3, 1], vertical_alignment="bottom")
    with c_title: st.title("團隊器材中心")
    with c_actions:
        if st.session_state.is_admin:
            b1, b2 = st.columns(2, gap="small")
            b1.button("➕ 新增", on_click=show_add_modal, use_container_width=True)
            b2.button("登出", on_click=perform_logout, type="primary", use_container_width=True)
        else:
            st.button("🔐 管理員登入", on_click=lambda: go_to("login"), type="primary", use_container_width=True)

    if not df.empty:
        # 計算總量與借出量
        total_items = len(df)
        total_qty = df['quantity'].sum()
        total_borrowed = df['borrowed'].sum()
        available_qty = total_qty - total_borrowed
        
        m1, m2, m3, m4 = st.columns(4)
        with m1: st.metric("📦 器材種類", total_items)
        with m2: st.metric("📊 庫存總數", int(total_qty))
        with m3: st.metric("✅ 剩餘可用", int(available_qty))
        with m4: st.metric("👤 目前借出", int(total_borrowed))

    st.write("")
    with st.container(border=True):
        search_query = st.text_input("🔍 搜尋器材...", placeholder="輸入關鍵字 (名稱、編號)...", label_visibility="collapsed")
        st.write("") 
        filter_options = ["全部顯示"] + CATEGORY_OPTIONS
        selected_category = st.pills("快速分類篩選", filter_options, default="全部顯示", label_visibility="collapsed")

    if not df.empty:
        if selected_category and selected_category != "全部顯示":
            filtered_df = df[df['category'] == selected_category]
        else:
            filtered_df = df

        if search_query:
            filtered_df = filtered_df[
                filtered_df['name'].str.contains(search_query, case=False) | 
                filtered_df['uid'].str.contains(search_query, case=False)
            ]
        
        if not filtered_df.empty:
            st.write("") 
            cols = st.columns(3)
            for i, (index, row) in enumerate(filtered_df.iterrows()):
                with cols[i % 3]:
                    with st.container(border=True):
                        img = row['image_url'] if row['image_url'] else "https://cdn-icons-png.flaticon.com/512/4992/4992482.png"
                        st.markdown(f'<div style="height:200px; overflow:hidden; border-radius:4px; display:flex; justify-content:center; background:#f0f2f6; margin-bottom:12px;"><img src="{img}" style="height:100%; width:100%; object-fit:cover;"></div>', unsafe_allow_html=True)
                        st.markdown(f"#### {row['name']}")
                        
                        # 🔥 動態計算狀態標籤
                        status_text, status_color = get_status_display(row)
                        st.caption(f"#{row['uid']} | 📍 {row['location']}")
                        st.markdown(f':{status_color}[**{status_text}**]')

                        st.markdown("---")
                        if st.session_state.is_admin:
                            if st.button("⚙️ 編輯 / 管理", key=f"btn_{row['uid']}", use_container_width=True):
                                show_edit_modal(row)
                        else:
                            # 🔥 借用邏輯：計算剩餘量
                            avail = row['quantity'] - row.get('borrowed', 0)
                            
                            # 如果是維修中或已借完，就不能勾選
                            is_disabled = (avail <= 0) or (row.get('status') in ['維修中', '報廢'])
                            
                            is_selected = row['uid'] in st.session_state.cart
                            
                            # 使用 key 綁定 checkbox
                            chk = st.checkbox(
                                "加入借用清單", 
                                key=f"check_{row['uid']}", 
                                value=is_selected,
                                disabled=is_disabled
                            )
                            
                            if chk and not is_selected:
                                st.session_state.cart[row['uid']] = 1 # 預設借 1 個
                                st.rerun()
                            elif not chk and is_selected:
                                del st.session_state.cart[row['uid']]
                                st.rerun()
                                    
        else:
            if selected_category != "全部顯示":
                st.info(f"📂 「{selected_category}」分類下目前沒有器材。")
            else:
                st.info("尚無符合搜尋條件的資料。")
    else: st.info("尚無資料")

# ==========================================
# 登入頁
# ==========================================
def login_page():
    render_header()
    _, c, _ = st.columns([1,5,1])
    with c:
        with st.container(border=True):
            st.markdown("<h2 style='text-align:center'>🔐 管理員登入</h2>", unsafe_allow_html=True)
            st.text_input("密碼", type="password", key="password_input")
            b1, b2 = st.columns(2)
            b1.button("取消", on_click=lambda: go_to("home"), use_container_width=True)
            b2.button("登入", type="primary", on_click=perform_login, use_container_width=True)

if st.session_state.current_page == "login": login_page()
else: main_page()
